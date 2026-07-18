"""DOCX text extraction module.

Preserves heading hierarchy and paragraph-level structure
for downstream formatting and analysis pipelines.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional


HEADING_STYLES = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
                  "Heading 4": 4, "Heading 5": 5, "Heading 6": 6}


@dataclass
class ParagraphData:
    """A single paragraph with its style and hierarchy information."""

    text: str
    style_name: str
    heading_level: int = 0  # 0 = body text, 1-6 = heading levels
    list_level: int = -1    # -1 = not a list item, 0+ = nested list depth
    is_list_item: bool = False


@dataclass
class DocxDocument:
    """Structured representation of an extracted DOCX document."""

    path: str
    paragraphs: list[ParagraphData] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        """Return text with heading markers (e.g., #, ##, ###)."""
        parts = []
        for para in self.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.heading_level > 0:
                prefix = "#" * para.heading_level
                parts.append(f"{prefix} {text}")
            elif para.is_list_item:
                indent = "  " * (para.list_level + 1) if para.list_level >= 0 else ""
                parts.append(f"{indent}- {text}")
            else:
                parts.append(text)
        return "\n\n".join(parts)

    @property
    def plain_text(self) -> str:
        """Return plain text without any markers."""
        return "\n\n".join(p.text.strip() for p in self.paragraphs if p.text.strip())

    @property
    def headings(self) -> list[ParagraphData]:
        """Return only heading paragraphs, useful for building a table of contents."""
        return [p for p in self.paragraphs if p.heading_level > 0]

    def get_section_text(self, heading_index: int = 0) -> str:
        """Get the full text of a section starting from a heading.

        Args:
            heading_index: Index into the headings list (0 = first heading).

        Returns:
            The heading plus all body text until the next heading at the same or
            higher level.
        """
        headings = self.headings
        if not headings or heading_index >= len(headings):
            return ""

        target = headings[heading_index]
        target_idx = self.paragraphs.index(target)

        # Find end boundary — next heading at same or higher (lower number) level
        end_idx = len(self.paragraphs)
        for i in range(target_idx + 1, len(self.paragraphs)):
            p = self.paragraphs[i]
            if p.heading_level > 0 and p.heading_level <= target.heading_level:
                end_idx = i
                break

        return "\n\n".join(
            p.text.strip()
            for p in self.paragraphs[target_idx:end_idx]
            if p.text.strip()
        )


def _resolve_heading_level(style_name: str) -> int:
    """Map python-docx style names to heading levels 1-6."""
    if style_name.startswith("Heading"):
        for i in range(1, 7):
            if f"Heading {i}" == style_name or style_name.endswith(f" {i}"):
                return i
    return 0


def parse_docx(path: str | Path) -> DocxDocument:
    """Parse a DOCX file and return structured text with heading hierarchy.

    Args:
        path: Path to the DOCX file.

    Returns:
        A DocxDocument containing paragraph-level structured data with
        heading levels, list information, and style metadata.

    Raises:
        FileNotFoundError: If the DOCX file does not exist.
        RuntimeError: If python-docx is not installed.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to parse DOCX files.") from exc

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    doc = Document(path)
    result = DocxDocument(path=str(path.resolve()))

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else "Normal"
        heading_level = _resolve_heading_level(style_name)

        # Detect list items
        is_list_item = False
        list_level = -1
        try:
            # python-docx list detection via numbering properties
            num_pr = paragraph._element.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
            )
            if num_pr is not None:
                is_list_item = True
                ilvl = num_pr.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl"
                )
                if ilvl is not None and ilvl.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                ) is not None:
                    list_level = int(ilvl.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                    ))
                else:
                    list_level = 0
        except Exception:
            pass

        result.paragraphs.append(
            ParagraphData(
                text=text,
                style_name=style_name,
                heading_level=heading_level,
                list_level=list_level,
                is_list_item=is_list_item,
            )
        )

    return result


def heading_tree(doc: DocxDocument) -> list[dict]:
    """Build a hierarchical tree from heading paragraphs.

    Returns:
        A nested list of dicts::
            [
                {"level": 1, "text": "Introduction", "children": [
                    {"level": 2, "text": "Background", "children": []}
                ]},
                ...
            ]
    """
    stack: list[dict] = []
    root: list[dict] = []

    for heading in doc.headings:
        node = {"level": heading.heading_level, "text": heading.text, "children": []}
        # Pop until we find a parent at a lower level
        while stack and stack[-1]["level"] >= heading.heading_level:
            stack.pop()
        if stack:
            stack[-1]["children"].append(node)
        else:
            root.append(node)
        stack.append(node)

    return root
