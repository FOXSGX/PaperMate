"""Formatting engine for Word documents.

Applies academic style profiles (GB/T 7713, IEEE, generic) to DOCX files,
adjusting fonts, sizes, spacing, heading numbering, page margins, and headers/footers.
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from . import styles


# ──────────────────────────────────────────────
#  Internal helpers
# ──────────────────────────────────────────────


def _pt_to_emu(pt: float) -> int:
    """Convert points to EMU (English Metric Units) used by OOXML."""
    return int(pt * 12700)


def _mm_to_emu(mm: float) -> int:
    """Convert millimeters to EMU."""
    return int(mm * 36000)


def _apply_font(run, font_spec: styles.FontSpec) -> None:
    """Apply font properties to a python-docx run."""
    from docx.shared import Pt, RGBColor

    run.font.name = font_spec.name

    # Handle East-Asian font fallback
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
    )
    if rfonts is None:
        from lxml import etree
        rfonts = etree.SubElement(
            rpr,
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts",
        )
    rfonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        font_spec.name,
    )

    run.font.size = Pt(font_spec.size_pt)
    run.bold = font_spec.bold
    run.italic = font_spec.italic
    if font_spec.color_rgb:
        run.font.color.rgb = RGBColor(
            int(font_spec.color_rgb[0:2], 16),
            int(font_spec.color_rgb[2:4], 16),
            int(font_spec.color_rgb[4:6], 16),
        )


def _set_alignment(paragraph, alignment: str):
    """Set paragraph alignment from a string value."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    paragraph.alignment = mapping.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)


def _heading_number(heading_level: int, numbering_format: str,
                    counters: list[int]) -> str:
    """Generate a heading number string from a format pattern.

    Digits in the format string are replaced by counter values in
    left-to-right order:

        counters = [2, 3], heading_level = 2
        numbering_format = "1.1"   → "2.3 "
        numbering_format = "第1章"  → "第2章 "
    """
    if not numbering_format:
        return ""

    result = list(numbering_format)
    digit_idx = 0
    for i, ch in enumerate(result):
        if ch.isdigit() and digit_idx < heading_level:
            if digit_idx < len(counters):
                result[i] = str(counters[digit_idx])
            digit_idx += 1

    return "".join(result) + " "


# ──────────────────────────────────────────────
#  Main formatting function
# ──────────────────────────────────────────────


def format_docx(
    input_path: str | Path,
    output_dir: str | Path,
    template: str = "gb7713",
    output_filename: Optional[str] = None,
) -> Path:
    """Apply academic formatting to a DOCX file.

    Args:
        input_path: Path to the input DOCX file.
        output_dir: Directory to write the formatted output.
        template: Style profile name (``gb7713``, ``generic``, ``ieee``).
        output_filename: Custom output filename (default:
                         ``{stem}_{template}_formatted.docx``).

    Returns:
        Path to the formatted output file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Emu
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise RuntimeError("python-docx is required for formatting.") from exc

    # ── Resolve style profile ──
    profile = styles.get_style(template)

    # ── Paths ──
    input_path = Path(input_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    if output_filename is None:
        output_filename = f"{input_path.stem}_{template}_formatted.docx"
    output_path = output_dir / output_filename

    # ── Open document ──
    document = Document(input_path)

    # ── Page margins ──
    section = document.sections[0]
    section.top_margin = Emu(_mm_to_emu(profile.page_margins.top_mm))
    section.bottom_margin = Emu(_mm_to_emu(profile.page_margins.bottom_mm))
    section.left_margin = Emu(_mm_to_emu(profile.page_margins.left_mm))
    section.right_margin = Emu(_mm_to_emu(profile.page_margins.right_mm))

    # ── Header / Footer ──
    if profile.header_text:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = profile.header_text
        if profile.header_font:
            for run in hp.runs:
                _apply_font(run, profile.header_font)

    if profile.footer_text:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.text = profile.footer_text
        if profile.footer_font:
            for run in fp.runs:
                _apply_font(run, profile.footer_font)

    # ── Iterate paragraphs and apply formatting ──
    counters: list[int] = []  # heading numbering counters

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else "Normal"
        heading_level = 0
        for i in range(1, 7):
            if f"Heading {i}" == style_name:
                heading_level = i
                break

        if heading_level > 0:
            _format_heading(paragraph, profile, heading_level, counters)
        else:
            _format_body(paragraph, profile)

    document.save(output_path)
    return output_path


def _format_heading(paragraph, profile: styles.StyleProfile,
                    level: int, counters: list[int]) -> None:
    """Apply heading-level formatting to a paragraph."""
    from docx.shared import Pt

    spec = profile.heading_for_level(level)

    # Update numbering counters
    while len(counters) < level:
        counters.append(0)
    counters[level - 1] += 1
    # Reset lower-level counters
    for i in range(level, len(counters)):
        counters[i] = 0

    # Paragraph spacing
    pf = paragraph.paragraph_format
    pf.space_before = Pt(spec.before_spacing_pt)
    pf.space_after = Pt(spec.after_spacing_pt)
    pf.line_spacing = 1.0  # headings are typically single-spaced
    _set_alignment(paragraph, spec.alignment)

    # Apply font to all runs
    for run in paragraph.runs:
        _apply_font(run, spec.font)

    # Prepend numbering if configured
    num_str = _heading_number(level, spec.numbering_format, counters)
    if num_str and not paragraph.text.strip().startswith(num_str.strip()):
        # Insert numbering at the start of the first run
        from docx.oxml import OxmlElement
        first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
        existing = first_run.text
        first_run.text = num_str + existing


def _format_body(paragraph, profile: styles.StyleProfile) -> None:
    """Apply body-text formatting to a paragraph."""
    from docx.shared import Pt

    pf = paragraph.paragraph_format
    pf.line_spacing = profile.body_line_spacing
    pf.first_line_indent = Pt(profile.body_first_line_indent_pt)
    _set_alignment(paragraph, profile.body_alignment)

    for run in paragraph.runs:
        _apply_font(run, profile.body_font)


def auto_detect_style(main_text: str) -> str:
    """Heuristic auto-detection of the most appropriate style profile.

    Scans the document text and picks the best matching profile.
    """
    from . import styles

    cn_chars = sum(1 for c in main_text if "\u4e00" <= c <= "\u9fff")
    en_chars = sum(1 for c in main_text if c.isascii() and c.isalpha())

    if cn_chars > en_chars:
        return "gb7713"

    # Check for IEEE-specific keywords (crude heuristic)
    lower = main_text.lower()
    ieee_hits = sum(lower.count(kw) for kw in
                    ["abstract", "index terms", "ieee", "doi"])
    if ieee_hits >= 2:
        return "ieee"

    return "generic"
